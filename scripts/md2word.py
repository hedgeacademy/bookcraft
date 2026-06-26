#!/usr/bin/env python3
"""md2word — Markdown to Word (.docx) converter with classic Chinese book layout.

Uses pandoc for the core conversion, then python-docx for post-processing.
Generates a reference.docx template on first run with professional Chinese typesetting.
"""

import argparse
import subprocess
import sys
import os
import re
import tomllib
from datetime import datetime
from pathlib import Path


def ensure_pandoc():
    """Verify pandoc is available."""
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        print("ERROR: pandoc not found. Install it first.", file=sys.stderr)
        sys.exit(1)


def ensure_python_docx():
    """Verify python-docx is available, install if needed."""
    try:
        import docx
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "python-docx",
             "-i", "https://pypi.tuna.tsinghua.edu.cn/simple",
             "--trusted-host", "pypi.tuna.tsinghua.edu.cn", "-q"],
            check=True
        )


def create_reference_docx(output_path, title="", subtitle="", author=""):
    """Create a reference .docx with classic Chinese book typography styles.

    Style specifications (classic Chinese book layout):
    - Title (Title): 小二, 黑体, bold, centered, 1.5x line spacing
    - Subtitle (Subtitle): 小三, 楷体, centered, 1.5x line spacing
    - Heading 1: 三号(16pt), 黑体, bold, centered
    - Heading 2: 四号(14pt), 黑体, bold
    - Heading 3: 小四(12pt), 黑体, bold
    - Normal (body): 五号(10.5pt), 宋体, first-line indent 2 chars
    - Page: 16开(18.4x26cm), margins top/bottom/left/right = 2.5/1.8/2/2cm
    - Grid approximation: 40 chars per line, 40 lines per page
    """
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement

    doc = Document()

    # ── Page setup: classic Chinese book layout (16开, 40x40 grid approximation) ──
    section = doc.sections[0]
    section.page_width = Cm(18.4)
    section.page_height = Cm(26)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.7)
    section.footer_distance = Cm(1.3)
    section.different_first_page_header_footer = True

    # ── Helper: apply a style ──
    def set_style(style_name, font_name, font_size_pt, bold=False,
                  alignment=None, line_spacing=1.5, space_before=0, space_after=0,
                  first_line_indent=None, font_name_east_asia=None, exact_line=False):
        style = doc.styles[style_name]
        pf = style.paragraph_format
        pf.line_spacing = line_spacing
        if exact_line:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)

        if alignment is not None:
            pf.alignment = alignment

        if first_line_indent is not None:
            pf.first_line_indent = first_line_indent

        font = style.font
        font.name = font_name
        font.size = Pt(font_size_pt)
        font.bold = bold
        font.color.rgb = RGBColor(0, 0, 0)

        ea_name = font_name_east_asia or font_name
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), ea_name)

    # ── Define styles ──
    # Title
    set_style('Title', '黑体', 18, bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
              space_before=72, space_after=12, font_name_east_asia='黑体')

    # Subtitle
    set_style('Subtitle', '楷体', 15, bold=False,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
              space_before=0, space_after=24, font_name_east_asia='楷体')

    # Heading 1 — centered (章标题居中)
    set_style('Heading 1', '黑体', 16, bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
              space_before=24, space_after=12, font_name_east_asia='黑体')

    # Heading 2
    set_style('Heading 2', '黑体', 14, bold=True,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
              space_before=18, space_after=6, font_name_east_asia='黑体')

    # Heading 3
    set_style('Heading 3', '黑体', 12, bold=True,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
              space_before=12, space_after=6, font_name_east_asia='黑体')

    # Normal (body text)
    set_style('Normal', '宋体', 10.5, bold=False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
              space_before=0, space_after=0,
              first_line_indent=Cm(0.74),  # ~2 Chinese chars
              font_name_east_asia='宋体')

    # Body Text — pandoc maps many paragraphs here, must match Normal
    set_style('Body Text', '宋体', 10.5, bold=False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
              space_before=0, space_after=0,
              first_line_indent=Cm(0.74),
              font_name_east_asia='宋体')

    # Add page break before each Heading 1 (chapter separation)
    h1_style = doc.styles['Heading 1']
    pPr = h1_style.element.get_or_add_pPr()
    page_break = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} />')
    pPr.append(page_break)

    # ── Front matter page (not needed in reference, just a placeholder) ──
    p = doc.add_paragraph("图书标题", style='Title')
    if subtitle:
        doc.add_paragraph(subtitle, style='Subtitle')
    if author:
        p = doc.add_paragraph(author, style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None

    doc.add_page_break()

    # Sample heading and body to demonstrate styles
    doc.add_heading("第一章  示例章节", level=1)
    doc.add_heading("1.1 示例小节", level=2)
    doc.add_paragraph(
        "这是一段示例正文，用于展示经典中文图书排版风格。"
        "正文使用宋体小四号字，首行缩进两个字符，1.5倍行距，"
        "两端对齐，A4页面。这是中文图书最通用、最经典的排版方式。"
    )
    doc.add_heading("1.2 另一个小节", level=2)
    doc.add_paragraph(
        "这是第二段示例正文。章节之间通过分页明确分隔，"
        "每个新章节从新页开始，便于阅读和印制。"
    )

    doc.save(output_path)
    print(f"Reference template created: {output_path}", file=sys.stderr)


def pandoc_convert(input_md, output_docx, reference_docx):
    """Run pandoc to convert Markdown to Word.

    --number-sections adds auto-numbering (1, 1.1, 1.1.1).
    H1 numbering is replaced with "第X章" in post-processing.
    TOC is a Word field that auto-updates on open (updateFields in settings.xml).
    """
    cmd = [
        "pandoc",
        input_md,
        "-o", output_docx,
        "--reference-doc", reference_docx,
        "--from", "markdown+smart",
        "--standalone",
        "--number-sections",
        "--metadata", "lang=zh-CN",
    ]
    print(f"Running: {' '.join(cmd)}", file=sys.stderr)
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"pandoc error:\n{result.stderr}", file=sys.stderr)
        sys.exit(result.returncode)
    if result.stderr:
        print(result.stderr, file=sys.stderr)



def _insert_copyright_page(doc, docx_path, title, subtitle, author):
    """Insert copyright page after title page. Modifies doc in place, then saves."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    config_dir = os.path.dirname(os.path.abspath(docx_path))
    config_candidates = [
        os.path.join(config_dir, 'book.toml'),
        os.path.join(os.getcwd(), 'book.toml'),
        os.path.join(os.path.dirname(config_dir), 'book.toml'),
    ]
    config_path = next((p for p in config_candidates if os.path.isfile(p)), config_candidates[0])

    cfg = {}
    if os.path.isfile(config_path):
        with open(config_path, 'rb') as f:
            cfg = tomllib.load(f)

    bookcraft = cfg.get('bookcraft', {})
    std_book = cfg.get('book', {})
    bookcraft_book = bookcraft.get('book', {})
    book = dict(bookcraft_book or {})
    if std_book:
        book.setdefault('title', std_book.get('title', title))
        authors = std_book.get('authors') or []
        if authors:
            book.setdefault('author', '、'.join(authors))
    pub = bookcraft.get('publisher', cfg.get('publisher', {}))
    ed = bookcraft.get('edition', cfg.get('edition', {}))
    cr = bookcraft.get('copyright_notice', cfg.get('copyright_notice', {}))
    desc = bookcraft.get('description', cfg.get('description', {}))
    cip = bookcraft.get('cip', cfg.get('cip', {}))

    body = doc.element.body

    def _add_para(text, font_size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=0, font_name='\u5b8b\u4f53',
                  first_line_indent=None, left_indent=None, right_indent=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.5
        if left_indent is not None:
            p.paragraph_format.left_indent = left_indent
        if right_indent is not None:
            p.paragraph_format.right_indent = right_indent
        if first_line_indent is not None:
            p.paragraph_format.first_line_indent = first_line_indent
        else:
            p.paragraph_format.first_line_indent = None
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = font_name
        run.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
        return p

    def _add_sep():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run('\u2014' * 40)
        run.font.size = Pt(7)
        return p

    # 扉页和版权页之间不换页，版权内容紧接扉页之后
    paras = []

    # ── 1. 内 容 简 介 (top, merged 简介+摘要) ──
    desc_text = desc.get('text', '').strip()
    if desc_text:
        paras.append(_add_para('内 容 简 介', font_size=10.5, bold=True, font_name='黑体',
                               align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=6))
        # Blank line after the heading for a compact copyright-page layout.
        paras.append(_add_para('', font_size=9, space_before=0, space_after=3))
        paras.append(_add_para(desc_text, font_size=9, first_line_indent=None,
                               left_indent=None, right_indent=None,
                               space_before=0, space_after=6))
        paras.append(_add_sep())


    # ── 2. CIP data ──
    cip_text = cip.get('text', '').strip()
    if cip_text:
        for line in cip_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            paras.append(_add_para(line, font_size=8 if '在版编目' in line else 7.5,
                                   bold=('在版编目' in line), space_before=0, space_after=0))
        paras.append(_add_sep())

    # ── 4. Book info ──
    bt = book.get('title', title)
    bs = book.get('subtitle', subtitle)
    ba = book.get('author', author)
    if bt:
        paras.append(_add_para(f'书名：{bt}', font_size=8, space_before=2))
    if bs:
        paras.append(_add_para(f'副书名：{bs}', font_size=8))
    if ba:
        paras.append(_add_para(f'作者：{ba}', font_size=8, space_after=2))

    # Copyright notice
    cr_text = cr.get('text', '').strip()
    if cr_text:
        paras.append(_add_para(cr_text, font_size=7.5, space_before=1))
        paras.append(_add_sep())
    else:
        paras.append(_add_sep())

    # ── 5. Publisher & edition ──
    pub_name = pub.get('name', '')
    if pub_name:
        paras.append(_add_para(f'出版发行：{pub_name}', font_size=7.5, space_before=2))
    pub_addr = pub.get('address', '')
    if pub_addr:
        paras.append(_add_para(f'地    址：{pub_addr}', font_size=7.5))
    pub_web = pub.get('website', '')
    if pub_web:
        paras.append(_add_para(f'网    址：{pub_web}', font_size=7.5, space_after=2))

    for key, label in [('version', '版    次'), ('printing', '印    次'),
                       ('format', '开    本'), ('isbn', '书    号'), ('price', '定    价')]:
        val = ed.get(key, '')
        if val:
            paras.append(_add_para(f'{label}：{val}', font_size=7.5))

    # ── 6. Disclaimer ──
    disc = cr.get('disclaimer', '').strip()
    if disc:
        paras.append(_add_sep())
        paras.append(_add_para(disc, font_size=7, space_before=1))

    # Section break after copyright page (the first front-matter item starts on a new page)
    tail_sect = doc.add_paragraph()
    pPr_tail = tail_sect.paragraph_format._element.get_or_add_pPr()
    sect_pr_tail = parse_xml(
        f'<w:sectPr {nsdecls("w")}>'
        f'<w:type w:val="nextPage"/>'
        f'<w:pgSz w:w="10431" w:h="14740"/>'
        f'<w:pgMar w:top="1417" w:right="1134" w:bottom="1021" w:left="1134" w:header="964" w:footer="737" w:gutter="0"/>'
        f'</w:sectPr>'
    )
    pPr_tail.append(sect_pr_tail)
    paras.append(tail_sect)

    # Insert all into body
    body_elem = doc.element.body
    insert_pos = 4  # after cover elements (0:title, 1:subtitle, 2:author, 3:sectPr)
    for p in reversed(paras):
        body_elem.remove(p._element)
        body_elem.insert(insert_pos, p._element)

    doc.save(docx_path)
    return doc

def post_process_docx(docx_path, title="", subtitle="", author=""):
    """Apply final polish to the generated .docx.

    Strategy: insert real section breaks before each Heading 1, then write
    chapter titles directly into per-section headers (no STYLEREF field needed).
    Footer gets a PAGE field for auto page numbering.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement
    from lxml import etree
    import copy
    import os
    import tempfile
    import zipfile
    import shutil

    # ── Pass 1: read the docx and identify Heading 1 positions ──
    doc = Document(docx_path)

    import re

    # Legacy projects may still concatenate Markdown title/copyright pages into
    # manuscript.md. Remove those leading blocks before generating the canonical
    # title and copyright pages from book.toml.
    duplicate_front_titles = {
        '扉页', '书名页', '版权页', '版权页与内容提要', '版权页 + 内容提要',
        '版权页+内容提要', '内容提要'
    }
    if title:
        duplicate_front_titles.add(title.strip())

    paragraphs = list(doc.paragraphs)
    remove_elements = []
    content_started = False
    index = 0
    while index < len(paragraphs):
        para = paragraphs[index]
        if para.style.name != 'Heading 1':
            index += 1
            continue
        heading = re.sub(r'^\d+\s+', '', para.text.strip())
        if not content_started and heading in duplicate_front_titles:
            end = index + 1
            while end < len(paragraphs) and paragraphs[end].style.name != 'Heading 1':
                end += 1
            remove_elements.extend(p._element for p in paragraphs[index:end])
            index = end
            continue
        content_started = True
        index += 1

    if remove_elements:
        for element in remove_elements:
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        doc.save(docx_path)
        doc = Document(docx_path)

    h1_data = []  # list of (para_index, title_text, original_h1_element)
    FRONT_MATTER_TITLES = {
        '推荐序', '序言', '序言一', '序言二', '序言三',
        '自序', '前言', '参考文献'
    }

    def _cn_to_int(value):
        value = value.strip()
        if value.isdigit():
            return int(value)
        nums = {'零': 0, '一': 1, '二': 2, '两': 2, '三': 3, '四': 4, '五': 5,
                '六': 6, '七': 7, '八': 8, '九': 9}
        if value == '十':
            return 10
        if value.startswith('十'):
            return 10 + nums.get(value[1:], 0)
        if '十' in value:
            left, right = value.split('十', 1)
            return nums.get(left, 0) * 10 + (nums.get(right, 0) if right else 0)
        return nums.get(value)

    def _normalize_chapter_title(title, chapter_no):
        title = title.strip()
        m = re.match(r'^第([一二两三四五六七八九十\d]+)章\s*(.*)$', title)
        if m:
            rest = m.group(2).strip()
            return f'第{chapter_no}章 {rest}' if rest else f'第{chapter_no}章'
        return title

    def _force_run_black(para):
        for run in para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    # ── Build pandoc_num → logical_chapter mapping ──
    pandoc_to_logical = {}  # pandoc chapter num → real chapter num
    logical_chapter = 0
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            text = para.text.strip()
            m = re.match(r'^(\d+)\s+(.+)', text)
            if m:
                pandoc_num = int(m.group(1))
                title_part = m.group(2)
                if title_part not in FRONT_MATTER_TITLES and not re.match(r'^参考文献$', title_part):
                    logical_chapter += 1
                    pandoc_to_logical[pandoc_num] = logical_chapter

    # ── Fix H1 titles: strip pandoc auto-number, keep original title ──
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 1':
            text = para.text.strip()
            m = re.match(r'^(\d+)\s+(.+)', text)
            if m:
                title_part = m.group(2)
                if title_part not in FRONT_MATTER_TITLES:
                    logical = pandoc_to_logical.get(int(m.group(1)))
                    if logical is not None:
                        title_part = _normalize_chapter_title(title_part, logical)
                if para.runs:
                    for run in para.runs:
                        run.text = ''
                    para.runs[0].text = title_part
                text = title_part
            # Ensure H1 is centered
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _force_run_black(para)
            h1_data.append((i, text, para))

    # ── Fix H2/H3 numbering using the nearest preceding real chapter ──
    # Pandoc counts front-matter H1 headings, so its raw H2 may be 2.1 even
    # when the visible book chapter is 第1章. Track the current real chapter
    # directly instead of trusting Pandoc's first numeric component.
    current_logical_chapter = 0
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            heading_text = para.text.strip()
            if heading_text not in FRONT_MATTER_TITLES and not re.match(r'^参考文献$', heading_text):
                current_logical_chapter += 1
            continue
        if para.style.name in ('Heading 2', 'Heading 3'):
            text = para.text.strip()
            # Match pandoc auto-number: "N.M Title" or "N.M.K Title"
            m = re.match(r'^(\d+)\.(\d+(?:\.\d+)*)\s+(.+)', text)
            if m:
                sub_number = m.group(2)
                title_part = m.group(3)
                # Strip any manual numbering already in title (e.g., "1.1 标题" → "标题")
                title_part = re.sub(r'^\d+(?:\.\d+)+\s+', '', title_part)
                if current_logical_chapter > 0:
                    new_text = f'{current_logical_chapter}.{sub_number} {title_part}'
                else:
                    # Front-matter subsections: strip number, keep title only
                    new_text = title_part
                if para.runs:
                    for run in para.runs:
                        run.text = ''
                    para.runs[0].text = new_text
                _force_run_black(para)

    # ── Remove generation labels and chapter-end progress markers ──
    def _replace_para_text(para, value):
        if para.runs:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = value
        else:
            para.add_run(value)

    for para in list(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        epigraph_with_text = re.match(r'^章前引文\s*[：:]\s*(.+)$', text)
        if epigraph_with_text:
            _replace_para_text(para, epigraph_with_text.group(1).strip())
            continue

        remove_para = (
            re.match(r'^章前引文\s*[：:]?\s*$', text)
            or re.match(
                r'^第\s*[一二两三四五六七八九十百零\d]+\s*章\s*'
                r'(?:完|结束)(?:\s*[，,、；;。.]?\s*本章.*)?$',
                text,
            )
            or re.match(r'^本章\s*(?:约|共|字数|预计).*(?:字|字符)?\s*[。.]?$', text)
        )
        if remove_para:
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)

    if not h1_data:
        # No chapters? Just do basic cleanup
        for para in doc.paragraphs:
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                para.paragraph_format.first_line_indent = None
        doc.save(docx_path)
        return doc

    # ── Remove first-line indent from centered, quotes, lists ──
    for para in doc.paragraphs:
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            para.paragraph_format.first_line_indent = None
        if para.style.name.startswith('Quote') or para.style.name.startswith('List'):
            para.paragraph_format.first_line_indent = None
        if para.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
            _force_run_black(para)

    # ── Fix numbered list spacing: reduce gap between number and text ──
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for para in doc.paragraphs:
        pPr = para._element.find(f'{{{_W}}}pPr')
        if pPr is not None and pPr.find(f'{{{_W}}}numPr') is not None:
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.first_line_indent = Cm(-0.5)

    doc.save(docx_path)

    # ── Pass 2: insert section breaks via raw XML manipulation ──
    # We need to insert <w:p><w:pPr><w:sectPr>...</w:sectPr></w:pPr></w:p>
    # before every Heading 1 paragraph except the first in the body.

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    tmp_path = docx_path + '.tmp'

    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    root = etree.fromstring(data)
                    body = root.find(f'{{{W}}}body')
                    if body is None:
                        zout.writestr(item, data)
                        continue

                    # Find all Heading 1 paragraphs in body
                    all_p = list(body)
                    h1_elements = []
                    for el in all_p:
                        if el.tag == f'{{{W}}}p':
                            pPr = el.find(f'{{{W}}}pPr')
                            if pPr is not None:
                                pStyle = pPr.find(f'{{{W}}}pStyle')
                                if pStyle is not None and pStyle.get(f'{{{W}}}val') == 'Heading1':
                                    h1_elements.append(el)

                    # ── Table caption & border pass ──
                    # Scan body for tables and insert captions "表X-Y 描述" above each
                    chapter_num = 0
                    table_count = 0
                    insertions = []  # (body_index, element_to_insert)
                    prev_para_text = ""  # track preceding paragraph for caption hint

                    for idx, el in enumerate(list(body)):
                        if el.tag == f'{{{W}}}p':
                            # Track preceding paragraph text for table caption description
                            texts = el.findall('.//' + f'{{{W}}}t')
                            para_text = ''.join(t.text or '' for t in texts).strip()
                            if para_text:
                                prev_para_text = para_text

                            pPr = el.find(f'{{{W}}}pPr')
                            if pPr is not None:
                                pStyle = pPr.find(f'{{{W}}}pStyle')
                                if pStyle is not None:
                                    style_val = pStyle.get(f'{{{W}}}val', '')

                                    # ── Fix H2/H3 tab spacing: replace <w:tab/> with single space ──
                                    if style_val in ('Heading2', 'Heading3'):
                                        for run in el.findall(f'{{{W}}}r'):
                                            tab = run.find(f'{{{W}}}tab')
                                            if tab is not None:
                                                run.remove(tab)
                                                space_t = etree.SubElement(run, f'{{{W}}}t')
                                                space_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                                                space_t.text = ' '
                                                break

                                    if style_val == 'Heading1':
                                        import re as re_mod
                                        h1_text = ''.join(t.text or '' for t in el.findall('.//' + f'{{{W}}}t'))
                                        m = re_mod.match(r'第([一二三四五六七八九十\d]+)章', h1_text)
                                        if m:
                                            cn = m.group(1)
                                            cn_map = {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':10}
                                            chapter_num = cn_map.get(cn, int(cn) if cn.isdigit() else 0)
                                        else:
                                            chapter_num = 0
                                        table_count = 0

                        elif el.tag == f'{{{W}}}tbl':
                            table_count += 1

                            # Extract table header (first row) for caption description
                            header_text = ""
                            first_row = el.find(f'{{{W}}}tr')
                            if first_row is not None:
                                cells = first_row.findall(f'{{{W}}}tc')
                                headers = []
                                for cell in cells[:4]:  # up to 4 columns
                                    cell_texts = cell.findall('.//' + f'{{{W}}}t')
                                    cell_text = ''.join(t.text or '' for t in cell_texts).strip()
                                    if cell_text:
                                        headers.append(cell_text)
                                if headers:
                                    header_text = '　' + '·'.join(headers)

                            if chapter_num > 0:
                                cap_text = f'表{chapter_num}-{table_count}{header_text}'
                            else:
                                cap_text = f'附表{table_count}{header_text}'

                            # Escape XML special chars in caption
                            cap_text = cap_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                            # Create caption paragraph
                            cap_para = etree.fromstring(f'''<w:p {nsdecls("w")}>
                              <w:pPr>
                                <w:jc w:val="center"/>
                                <w:spacing w:before="120" w:after="60"/>
                              </w:pPr>
                              <w:r>
                                <w:rPr>
                                  <w:rFonts w:eastAsia="黑体" w:ascii="黑体" w:hAnsi="黑体"/>
                                  <w:b/>
                                  <w:sz w:val="18"/>
                                </w:rPr>
                                <w:t>{cap_text}</w:t>
                              </w:r>
                            </w:p>''')
                            insertions.append((idx, cap_para))

                            # ── Table borders: open left/right (三线表变体) ──
                            tblPr = el.find(f'{{{W}}}tblPr')
                            if tblPr is None:
                                tblPr = etree.SubElement(el, f'{{{W}}}tblPr')
                                el.insert(0, tblPr)

                            # ── Table indent: match body text first-line indent (~2 Chinese chars) ──
                            old_ind = tblPr.find(f'{{{W}}}tblInd')
                            if old_ind is not None:
                                tblPr.remove(old_ind)
                            tblInd = etree.SubElement(tblPr, f'{{{W}}}tblInd')
                            tblInd.set(f'{{{W}}}w', '420')  # ~0.74cm
                            tblInd.set(f'{{{W}}}type', 'dxa')

                            # Remove old borders if any
                            old_borders = tblPr.find(f'{{{W}}}tblBorders')
                            if old_borders is not None:
                                tblPr.remove(old_borders)

                            tblBorders = etree.SubElement(tblPr, f'{{{W}}}tblBorders')
                            # Top border (thick)
                            top = etree.SubElement(tblBorders, f'{{{W}}}top')
                            top.set(f'{{{W}}}val', 'single')
                            top.set(f'{{{W}}}sz', '12')
                            top.set(f'{{{W}}}space', '0')
                            top.set(f'{{{W}}}color', '000000')
                            # Bottom border (thick)
                            bottom = etree.SubElement(tblBorders, f'{{{W}}}bottom')
                            bottom.set(f'{{{W}}}val', 'single')
                            bottom.set(f'{{{W}}}sz', '12')
                            bottom.set(f'{{{W}}}space', '0')
                            bottom.set(f'{{{W}}}color', '000000')
                            # Inside horizontal (thin, separates header from body)
                            insideH = etree.SubElement(tblBorders, f'{{{W}}}insideH')
                            insideH.set(f'{{{W}}}val', 'single')
                            insideH.set(f'{{{W}}}sz', '4')
                            insideH.set(f'{{{W}}}space', '0')
                            insideH.set(f'{{{W}}}color', '000000')
                            # Inside vertical (thin)
                            insideV = etree.SubElement(tblBorders, f'{{{W}}}insideV')
                            insideV.set(f'{{{W}}}val', 'single')
                            insideV.set(f'{{{W}}}sz', '4')
                            insideV.set(f'{{{W}}}space', '0')
                            insideV.set(f'{{{W}}}color', 'CCCCCC')
                            # Left & right: open (no border)
                            left = etree.SubElement(tblBorders, f'{{{W}}}left')
                            left.set(f'{{{W}}}val', 'none')
                            left.set(f'{{{W}}}sz', '0')
                            left.set(f'{{{W}}}space', '0')
                            right = etree.SubElement(tblBorders, f'{{{W}}}right')
                            right.set(f'{{{W}}}val', 'none')
                            right.set(f'{{{W}}}sz', '0')
                            right.set(f'{{{W}}}space', '0')

                            # ── Gray shading on first row (header) ──
                            first_row = el.find(f'{{{W}}}tr')
                            if first_row is not None:
                                trPr = first_row.find(f'{{{W}}}trPr')
                                if trPr is None:
                                    trPr = etree.SubElement(first_row, f'{{{W}}}trPr')
                                    first_row.insert(0, trPr)
                                # Add shading to each cell in header row
                                for tc in first_row.findall(f'{{{W}}}tc'):
                                    tcPr = tc.find(f'{{{W}}}tcPr')
                                    if tcPr is None:
                                        tcPr = etree.SubElement(tc, f'{{{W}}}tcPr')
                                        tc.insert(0, tcPr)
                                    # Remove old shading
                                    old_shd = tcPr.find(f'{{{W}}}shd')
                                    if old_shd is not None:
                                        tcPr.remove(old_shd)
                                    shd = etree.SubElement(tcPr, f'{{{W}}}shd')
                                    shd.set(f'{{{W}}}val', 'clear')
                                    shd.set(f'{{{W}}}color', 'auto')
                                    shd.set(f'{{{W}}}fill', 'E7E6E6')

                    # Apply insertions (reverse order to preserve indices)
                    for idx, cap_para in reversed(insertions):
                        body.insert(idx, cap_para)

                    # ── Section breaks: each H1 starts a separate section ──
                    FRONT_MATTER = set()
                    section_count = 1
                    skipped_first = False
                    for h1_el in h1_elements:
                        h1_text = ''.join(t.text or '' for t in h1_el.findall('.//' + f'{{{W}}}t')).strip()
                        if not skipped_first:
                            skipped_first = True
                            continue  # always skip the first H1
                        if h1_text in FRONT_MATTER:
                            continue  # front matter, no section break
                        # Create section-break paragraph
                        sect_pr = etree.SubElement(
                            etree.Element(f'{{{W}}}pPr'),
                            f'{{{W}}}sectPr'
                        )
                        sect_pr.set(f'{{{R}}}id', f'rId{100 + section_count}')
                        sect_type = etree.SubElement(sect_pr, f'{{{W}}}type')
                        sect_type.set(f'{{{W}}}val', 'nextPage')
                        pgSz = etree.SubElement(sect_pr, f'{{{W}}}pgSz')
                        pgSz.set(f'{{{W}}}w', '10431')
                        pgSz.set(f'{{{W}}}h', '14740')
                        pgMar = etree.SubElement(sect_pr, f'{{{W}}}pgMar')
                        pgMar.set(f'{{{W}}}top', '1417')
                        pgMar.set(f'{{{W}}}right', '1134')
                        pgMar.set(f'{{{W}}}bottom', '1021')
                        pgMar.set(f'{{{W}}}left', '1134')
                        pgMar.set(f'{{{W}}}header', '964')
                        pgMar.set(f'{{{W}}}footer', '737')
                        pgMar.set(f'{{{W}}}gutter', '0')

                        pPr_wrapper = etree.Element(f'{{{W}}}pPr')
                        pPr_wrapper.append(sect_pr)
                        break_para = etree.Element(f'{{{W}}}p')
                        break_para.append(pPr_wrapper)

                        body_index = list(body).index(h1_el)
                        body.insert(body_index, break_para)
                        section_count += 1

                    data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

                elif item.filename == 'word/settings.xml':
                    # Add updateFields for PAGE fields
                    root = etree.fromstring(data)
                    existing = root.find(f'{{{W}}}updateFields')
                    if existing is None:
                        uf = etree.SubElement(root, f'{{{W}}}updateFields')
                        uf.set(f'{{{W}}}val', 'true')
                    else:
                        existing.set(f'{{{W}}}val', 'true')
                    data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

                zout.writestr(item, data)

    # Replace original
    os.replace(tmp_path, docx_path)

    # ── Pass 2.4: Add cover page ──
    doc_cover = Document(docx_path)
    body_elem = doc_cover.element.body

    # Cover section break (next page)
    cover_sect = doc_cover.add_paragraph()
    pPr = cover_sect.paragraph_format._element.get_or_add_pPr()
    sect_pr = parse_xml(
        f'<w:sectPr {nsdecls("w")}>'
        f'<w:type w:val="nextPage"/>'
        f'<w:pgSz w:w="10431" w:h="14740"/>'
        f'<w:pgMar w:top="1417" w:right="1134" w:bottom="1021" w:left="1134" w:header="964" w:footer="737" w:gutter="0"/>'
        f'</w:sectPr>'
    )
    pPr.append(sect_pr)
    body_elem.remove(cover_sect._element)
    body_elem.insert(0, cover_sect._element)

    # Cover title
    cover_title = doc_cover.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.paragraph_format.space_before = Pt(180)
    run_t = cover_title.add_run(title if title else '')
    run_t.font.size = Pt(22)
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run_t._element.get_or_add_rPr()
    rPr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="黑体"/>'))
    body_elem.remove(cover_title._element)
    body_elem.insert(0, cover_title._element)

    # Cover subtitle
    cover_sub = doc_cover.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_sub.paragraph_format.space_before = Pt(24)
    run_s = cover_sub.add_run(subtitle if subtitle else '')
    run_s.font.size = Pt(16)
    run_s.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run_s._element.get_or_add_rPr()
    rPr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="楷体"/>'))
    body_elem.remove(cover_sub._element)
    body_elem.insert(1, cover_sub._element)

    # Cover author
    cover_auth = doc_cover.add_paragraph()
    cover_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_auth.paragraph_format.space_before = Pt(120)
    run_a = cover_auth.add_run(author if author else '')
    run_a.font.size = Pt(14)
    run_a.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run_a._element.get_or_add_rPr()
    rPr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体"/>'))
    body_elem.remove(cover_auth._element)
    body_elem.insert(2, cover_auth._element)

    doc_cover.save(docx_path)

    # ── Pass 2.45: Insert copyright page after title page ──
    doc_cover = _insert_copyright_page(doc_cover, docx_path, title, subtitle, author)

    # ── Pass 2.5: Build static TOC ──
    doc_toc = Document(docx_path)
    toc_entries = []
    _FRONT_H1 = {'推荐序', '序言', '序言一', '序言二', '序言三', '自序', '前言'}
    in_chapter = False  # only collect headings from first real chapter onward
    for para in doc_toc.paragraphs:
        if para.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
            h_text = para.text.strip()
            if not h_text:
                continue
            if para.style.name == 'Heading 1':
                if h_text in _FRONT_H1:
                    continue
                in_chapter = True  # first non-front H1 unlocks collection
            if not in_chapter:
                continue
            level = int(para.style.name.split()[-1])
            toc_entries.append((level, h_text))
    
    if toc_entries:
        # Insert TOC title
        toc_title = doc_toc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = toc_title.add_run('目　录')
        run_title.font.size = Pt(14)
        run_title.bold = True
        rPr = run_title._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="黑体"/>')
        rPr.insert(0, rFonts)
        
        # Move TOC title after cover (4 elements: title, subtitle, author, section break)
        body_elem = doc_toc.element.body
        # Find TOC insertion point: after 前言, before first chapter H1
        _FRONT_MATTER_TOC = {
            '推荐序', '序言', '序言一', '序言二', '序言三', '自序', '前言'
        }
        toc_start = 4  # fallback
        for idx, el in enumerate(list(body_elem)):
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag == 'p':
                pPr_el = el.find(f'{{{W}}}pPr')
                if pPr_el is not None:
                    pStyle_el = pPr_el.find(f'{{{W}}}pStyle')
                    if pStyle_el is not None and pStyle_el.get(f'{{{W}}}val') == 'Heading1':
                        h1_txt = ''.join(t.text or '' for t in el.findall('.//' + f'{{{W}}}t')).strip()
                        if h1_txt not in _FRONT_MATTER_TOC:
                            toc_start = idx
                            break
        body_elem.remove(toc_title._element)
        body_elem.insert(toc_start, toc_title._element)
        
        # Insert TOC entries after title
        for level, h_text in reversed(toc_entries):
            indent = Cm((level - 1) * 0.8)
            font_sz = Pt(10.5) if level == 1 else Pt(10)
            
            toc_para = doc_toc.add_paragraph()
            toc_para.paragraph_format.left_indent = indent
            toc_para.paragraph_format.first_line_indent = None
            toc_para.paragraph_format.space_before = Pt(2)
            toc_para.paragraph_format.space_after = Pt(2)
            
            run = toc_para.add_run(h_text)
            run.font.size = font_sz
            run.bold = (level == 1)
            rPr = run._element.get_or_add_rPr()
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>')
            rPr.insert(0, rFonts)
            
            # Move before the content (after cover + TOC title)
            body_elem.remove(toc_para._element)
            body_elem.insert(toc_start + 1, toc_para._element)
        
        # Section break after TOC (makes TOC its own section, no header)
        pb = doc_toc.add_paragraph()
        pPr = pb.paragraph_format._element.get_or_add_pPr()
        sect_pr = parse_xml(
        f'<w:sectPr {nsdecls("w")}>'
        f'<w:type w:val="nextPage"/>'
        f'<w:pgSz w:w="10431" w:h="14740"/>'
        f'<w:pgMar w:top="1417" w:right="1134" w:bottom="1021" w:left="1134" w:header="964" w:footer="737" w:gutter="0"/>'
        f'</w:sectPr>'
        )
        pPr.append(sect_pr)
        body_elem.remove(pb._element)
        body_elem.insert(toc_start + 1 + len(toc_entries), pb._element)
    
    doc_toc.save(docx_path)

    # ── Pass 2.6: Format the single required 自序 signature block ──
    doc_pro = Document(docx_path)
    in_author_preface = False
    author_preface_paras = []
    for para in doc_pro.paragraphs:
        if para.style.name == 'Heading 1' and para.text.strip() == '自序':
            in_author_preface = True
            continue
        if in_author_preface:
            if para.style.name == 'Heading 1':
                break
            author_preface_paras.append(para)
    if author_preface_paras:
        for para in author_preface_paras:
            if para.text.strip():
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        non_empty = [p for p in author_preface_paras if p.text.strip()]
        if len(non_empty) >= 2:
            signature_paras = non_empty[-2:]
            for para in signature_paras:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.first_line_indent = None

            first_signature = signature_paras[0]
            first_index = author_preface_paras.index(first_signature)
            has_blank_before = (
                first_index > 0 and not author_preface_paras[first_index - 1].text.strip()
            )
            if not has_blank_before:
                first_signature._p.addprevious(OxmlElement('w:p'))
    doc_pro.save(docx_path)

    # ── Pass 3: set per-section headers by scanning section content ──
    doc2 = Document(docx_path)
    sections = doc2.sections

    # Build paragraph→section mapping
    all_paras = doc2.paragraphs
    body_elems = list(doc2.element.body)
    current_sec = 0
    para_idx = 0
    para_to_sec = {}
    sec_first_para = {}
    for el in body_elems:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'p':
            para_to_sec[para_idx] = current_sec
            if current_sec not in sec_first_para:
                sec_first_para[current_sec] = para_idx
            para_idx += 1
            pPr = el.find(f'{{{W}}}pPr')
            if pPr is not None and pPr.find(f'{{{W}}}sectPr') is not None:
                current_sec += 1

    first_real_chapter = True  # first正文 section restarts with Arabic page 1.

    for sec_idx, section in enumerate(sections):
        section.different_first_page_header_footer = False
        header = section.header
        header.is_linked_to_previous = False

        # Find first H1 in this section
        chapter_title = ""
        start_p = sec_first_para.get(sec_idx, 0)
        section_has_toc = False
        for pi in range(start_p, min(start_p + 100, len(all_paras))):
            if para_to_sec.get(pi) == sec_idx and all_paras[pi].text.strip() == '目　录':
                section_has_toc = True
            if para_to_sec.get(pi) == sec_idx and all_paras[pi].style.name == 'Heading 1':
                chapter_title = all_paras[pi].text.strip()
                break

        is_front_matter = chapter_title in (
            '推荐序', '序言', '序言一', '序言二', '序言三', '自序', '前言'
        )
        is_toc = section_has_toc
        is_cover_or_copyright = chapter_title == '' and not section_has_toc
        is_body = not is_front_matter and not is_toc and not is_cover_or_copyright

        # No running header for cover/copyright/front-matter/TOC
        if not is_body:
            chapter_title = ""

        # ── Header paragraph ──
        # Clear default empty paragraph
        for p in header.paragraphs:
            p.clear()

        hp = header.paragraphs[0]
        if chapter_title:
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hp.paragraph_format.space_after = Pt(0)
            hp.paragraph_format.space_before = Pt(0)

            # Add chapter title as plain text
            run = hp.add_run(chapter_title)
            run.font.size = Pt(9)
            run.font.name = '宋体'
            rPr = run._element.get_or_add_rPr()
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体" w:ascii="宋体" w:hAnsi="宋体"/>')
            rPr.insert(0, rFonts)

            # Thin line separator under header
            pPr = hp._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

        # ── Footer: centered page number (only for chapters, starting at 1) ──
        footer = section.footer
        footer.is_linked_to_previous = False

        # Add footer paragraph if needed
        if not footer.paragraphs:
            footer.add_paragraph()
        for p in footer.paragraphs:
            p.clear()

        fp = footer.paragraphs[0]

        if is_cover_or_copyright or is_front_matter:
            # 扉页、版权页、序言一至三、自序和前言：无页码
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dummy_run = fp.add_run()
            dummy_run.text = ''
        else:
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = fp.add_run()
            run.font.size = Pt(9)

            fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._element.append(fld_begin)

            instr_run = fp.add_run()
            field_code = ' PAGE \\* ROMAN ' if is_toc else ' PAGE \\# "000" '
            instr_run._element.append(
                parse_xml(
                    f'<w:instrText {nsdecls("w")} xml:space="preserve">'
                    f'{field_code}</w:instrText>'
                )
            )

            fld_end_run = fp.add_run()
            fld_end_run._element.append(
                parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            )

            sect_pr = section._sectPr
            if is_toc:
                # 目录：大写罗马页码，从 I 开始。
                pgNumType = parse_xml(
                    f'<w:pgNumType {nsdecls("w")} w:fmt="upperRoman" w:start="1"/>'
                )
                sect_pr.append(pgNumType)
            elif is_body and first_real_chapter:
                # 正文：第一章从 1 重新开始，PAGE 字段按三位数字显示为 001。
                sect_pr = section._sectPr
                pgNumType = parse_xml(
                    f'<w:pgNumType {nsdecls("w")} w:fmt="decimal" w:start="1"/>'
                )
                sect_pr.append(pgNumType)
                first_real_chapter = False

    doc2.save(docx_path)
    print(f"  Sections: {len(doc2.sections)}, Chapters: {len(h1_data)}", file=sys.stderr)


def update_toc_page_numbers_with_word(docx_path):
    """Use Microsoft Word pagination to add page numbers to the static TOC.

    The static TOC is intentionally kept because it excludes front-matter
    headings. Page numbers require a layout engine, so this optional Windows
    pass uses Word COM when available. Non-Windows or headless environments
    keep the title-only static TOC.
    """
    if os.name != 'nt':
        return False

    import json
    import tempfile

    ps_script = r'''
param([string]$DocxPath)
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
try {
  $doc = $word.Documents.Open($DocxPath, $false, $false, $false)
  for ($pass = 1; $pass -le 2; $pass++) {
    $doc.Repaginate()
    $tocTitlePara = $null
    foreach ($p in $doc.Paragraphs) {
      $text = ($p.Range.Text.Trim() -replace '[\r\a]+$','')
      if ($text -eq '目　录') { $tocTitlePara = $p; break }
    }
    if (-not $tocTitlePara) { throw 'TOC title not found' }

    $firstChapterPara = $null
    foreach ($p in $doc.Paragraphs) {
      $text = ($p.Range.Text.Trim() -replace '[\r\a]+$','')
      $style = $p.Range.Style.NameLocal
      if (($style -match '^标题\s*1$|^Heading\s*1$') -and
          $text -ne '推荐序' -and $text -ne '序言' -and
          $text -ne '序言一' -and $text -ne '序言二' -and $text -ne '序言三' -and
          $text -ne '自序' -and $text -ne '前言' -and
          $p.Range.Start -gt $tocTitlePara.Range.End) {
        $firstChapterPara = $p
        break
      }
    }
    if (-not $firstChapterPara) { throw 'First chapter heading not found' }

    $entries = @()
    foreach ($p in $doc.Paragraphs) {
      if ($p.Range.Start -lt $firstChapterPara.Range.Start) { continue }
      $text = ($p.Range.Text.Trim() -replace '[\r\a]+$','')
      if ([string]::IsNullOrWhiteSpace($text)) { continue }
      $style = $p.Range.Style.NameLocal
      if ($style -match '^标题\s*([123])$|^Heading\s*([123])$') {
        $level = if ($Matches[1]) { [int]$Matches[1] } else { [int]$Matches[2] }
        # Use adjusted page number so section restarts are reflected in TOC entries.
        $page = $p.Range.Information(1)
        $entries += [PSCustomObject]@{ Level=$level; Text=$text; Page=$page }
      }
    }

    $start = $tocTitlePara.Range.End
    $end = $firstChapterPara.Range.Start
    $lines = New-Object System.Collections.Generic.List[string]
    foreach ($e in $entries) {
      $indent = if ($e.Level -eq 1) { '' } elseif ($e.Level -eq 2) { '    ' } else { '        ' }
      $pageText = "{0:D3}" -f [int]$e.Page
      $lines.Add($indent + $e.Text + "`t" + $pageText)
    }
    $replace = [string]::Join("`r", $lines) + "`r"
    $doc.Range($start, $end).Text = $replace
    # Replacing the TOC region removes the section break before the first
    # chapter. Put it back so front-matter and body page numbering stay separate.
    $breakRange = $doc.Range($start + $replace.Length, $start + $replace.Length)
    $breakRange.InsertBreak(2)

    $inserted = $doc.Range($start, $start + $replace.Length)
    foreach ($p in $inserted.Paragraphs) {
      $t = ($p.Range.Text.Trim() -replace '[\r\a]+$','')
      if ([string]::IsNullOrWhiteSpace($t)) { continue }
      $p.Range.Style = -1
      $p.Range.Font.NameFarEast = '宋体'
      $p.Range.Font.Name = 'Times New Roman'
      $p.Range.Font.Size = 10.5
      $p.Range.Font.Color = 0
      $p.Format.SpaceBefore = 2
      $p.Format.SpaceAfter = 2
      $p.Format.FirstLineIndent = 0
      $p.Format.LeftIndent = 0
      if ($t -match '^\s{8}') { $p.Format.LeftIndent = $word.CentimetersToPoints(1.6) }
      elseif ($t -match '^\s{4}') { $p.Format.LeftIndent = $word.CentimetersToPoints(0.8) }
      $p.Format.TabStops.ClearAll()
      [void]$p.Format.TabStops.Add($word.CentimetersToPoints(14.2), 2, 1)
    }
  }

  $doc.Repaginate()

  $tocTitlePara = $null
  foreach ($p in $doc.Paragraphs) {
    $text = ($p.Range.Text.Trim() -replace '[\r\a]+$','')
    if ($text -eq '目　录') { $tocTitlePara = $p; break }
  }
  $firstChapterPara = $null
  foreach ($p in $doc.Paragraphs) {
    $text = ($p.Range.Text.Trim() -replace '[\r\a]+$','')
    $style = $p.Range.Style.NameLocal
    if (($style -match '^标题\s*1$|^Heading\s*1$') -and
        $text -ne '推荐序' -and $text -ne '序言' -and
        $text -ne '序言一' -and $text -ne '序言二' -and $text -ne '序言三' -and
        $text -ne '自序' -and $text -ne '前言' -and
        $p.Range.Start -gt $tocTitlePara.Range.End) {
      $firstChapterPara = $p
      break
    }
  }
  if ($tocTitlePara -and $firstChapterPara) {
    $tocSec = $tocTitlePara.Range.Sections.Item(1)
    $tocSec.Footers.Item(1).LinkToPrevious = $false
    $tocPages = $tocSec.Footers.Item(1).PageNumbers
    $tocPages.RestartNumberingAtSection = $true
    $tocPages.StartingNumber = 1
    $tocPages.NumberStyle = 1
    $tocFooter = $tocSec.Footers.Item(1).Range
    $tocFooter.Text = ''
    $tocFooter.ParagraphFormat.Alignment = 1
    $tocField = $doc.Fields.Add($tocFooter, 33, '', $true)
    $tocField.Code.Text = ' PAGE \* ROMAN '
    [void]$tocField.Update()

    $bodySec = $firstChapterPara.Range.Sections.Item(1)
    $bodySec.Footers.Item(1).LinkToPrevious = $false
    $bodyPages = $bodySec.Footers.Item(1).PageNumbers
    $bodyPages.RestartNumberingAtSection = $true
    $bodyPages.StartingNumber = 1
    $bodyPages.NumberStyle = 0
    $bodyFooter = $bodySec.Footers.Item(1).Range
    $bodyFooter.Text = ''
    $bodyFooter.ParagraphFormat.Alignment = 1
    $bodyField = $doc.Fields.Add($bodyFooter, 33, '', $true)
    $bodyField.Code.Text = ' PAGE \# "000" '
    [void]$bodyField.Update()
  }
  foreach ($p in $doc.Paragraphs) {
    $text = ($p.Range.Text.Trim() -replace '[\r\a]+$','')
    $style = $p.Range.Style.NameLocal
    if ([string]::IsNullOrWhiteSpace($text) -and
        ($style -match '^标题\s*1$|^Heading\s*1$')) {
      $p.Range.Style = -1
    }
  }
  $doc.Save()
  $doc.Close($false)
} finally {
  $word.Quit()
}
'''

    with tempfile.NamedTemporaryFile('w', suffix='.ps1', delete=False, encoding='utf-8-sig') as f:
        f.write(ps_script)
        ps_path = f.name

    try:
        result = subprocess.run(
            ['powershell', '-NoProfile', '-ExecutionPolicy', 'Bypass',
             '-File', ps_path, os.path.abspath(docx_path)],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='replace',
            timeout=120,
        )
        if result.returncode != 0:
            print(f"WARNING: could not add TOC page numbers with Word COM:\n{result.stderr}",
                  file=sys.stderr)
            return False
        patch_toc_roman_page_numbering(docx_path)
        print("  TOC page numbers updated with Word pagination", file=sys.stderr)
        return True
    except Exception as exc:
        print(f"WARNING: could not add TOC page numbers with Word COM: {exc}", file=sys.stderr)
        return False
    finally:
        try:
            os.remove(ps_path)
        except OSError:
            pass


def patch_toc_roman_page_numbering(docx_path):
    """Force TOC Roman numbering and three-digit body page numbers.

    Word COM sometimes keeps a plain PAGE field even after setting
    PageNumbers.NumberStyle. Patch the OOXML directly so the TOC footer uses
    PAGE \\* ROMAN and the body section uses PAGE \\# "000".
    """
    from lxml import etree
    import zipfile
    import tempfile
    import shutil

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    PR = 'http://schemas.openxmlformats.org/package/2006/relationships'

    tmp_path = docx_path + '.tmp'

    with zipfile.ZipFile(docx_path, 'r') as zin:
        document_xml = zin.read('word/document.xml')
        rels_xml = zin.read('word/_rels/document.xml.rels')
        content_types_xml = zin.read('[Content_Types].xml')
        root = etree.fromstring(document_xml)
        rels_root = etree.fromstring(rels_xml)
        content_types_root = etree.fromstring(content_types_xml)
        relmap = {
            rel.get('Id'): rel.get('Target')
            for rel in rels_root.findall(f'{{{PR}}}Relationship')
        }
        new_parts = {}

        body = root.find(f'{{{W}}}body')
        section_index = 1
        toc_section = None
        body_section = None
        seen_toc = False
        section_props = []

        for el in list(body):
            if el.tag == f'{{{W}}}p':
                text = ''.join(t.text or '' for t in el.findall('.//' + f'{{{W}}}t')).strip()
                p_style = None
                pPr = el.find(f'{{{W}}}pPr')
                if pPr is not None:
                    pStyle = pPr.find(f'{{{W}}}pStyle')
                    if pStyle is not None:
                        p_style = pStyle.get(f'{{{W}}}val')
                if text == '目　录':
                    toc_section = section_index
                    seen_toc = True
                if (seen_toc and body_section is None and p_style in ('Heading1', '1')
                        and text and text not in (
                            '推荐序', '序言', '序言一', '序言二', '序言三', '自序', '前言'
                        )):
                    body_section = section_index
                if pPr is not None:
                    sectPr = pPr.find(f'{{{W}}}sectPr')
                    if sectPr is not None:
                        section_props.append(sectPr)
                        section_index += 1
        final_sect = body.find(f'{{{W}}}sectPr')
        if final_sect is not None:
            section_props.append(final_sect)

        def ensure_pgnum(sectPr, fmt, start=None):
            pg = sectPr.find(f'{{{W}}}pgNumType')
            if pg is None:
                pg = etree.SubElement(sectPr, f'{{{W}}}pgNumType')
            pg.set(f'{{{W}}}fmt', fmt)
            if start is not None:
                pg.set(f'{{{W}}}start', str(start))
            return pg

        def default_footer_target(sectPr):
            for fr in sectPr.findall(f'{{{W}}}footerReference'):
                if fr.get(f'{{{W}}}type') == 'default':
                    return fr, relmap.get(fr.get(f'{{{R}}}id'))
            return None, None

        def next_footer_name():
            nums = []
            for name in zin.namelist():
                if name.startswith('word/footer') and name.endswith('.xml'):
                    stem = name.rsplit('/', 1)[-1][6:-4]
                    if stem.isdigit():
                        nums.append(int(stem))
            for name in new_parts:
                if name.startswith('word/footer') and name.endswith('.xml'):
                    stem = name.rsplit('/', 1)[-1][6:-4]
                    if stem.isdigit():
                        nums.append(int(stem))
            return f'footer{(max(nums) if nums else 0) + 1}.xml'

        def next_rel_id():
            nums = []
            for rid in relmap:
                if rid.startswith('rId') and rid[3:].isdigit():
                    nums.append(int(rid[3:]))
            return f'rId{(max(nums) if nums else 0) + 1}'

        def make_footer_xml(instr):
            return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:p>
    <w:pPr><w:jc w:val="center"/></w:pPr>
    <w:r><w:fldChar w:fldCharType="begin"/></w:r>
    <w:r><w:instrText xml:space="preserve">{instr}</w:instrText></w:r>
    <w:r><w:fldChar w:fldCharType="separate"/></w:r>
    <w:r><w:t>{'I' if 'ROMAN' in instr else '001'}</w:t></w:r>
    <w:r><w:fldChar w:fldCharType="end"/></w:r>
  </w:p>
</w:ftr>'''.encode('utf-8')

        def add_footer_part(sectPr, instr):
            footer_name = next_footer_name()
            rid = next_rel_id()
            relmap[rid] = footer_name
            new_parts[f'word/{footer_name}'] = make_footer_xml(instr)

            rel = etree.SubElement(rels_root, f'{{{PR}}}Relationship')
            rel.set('Id', rid)
            rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer')
            rel.set('Target', footer_name)

            CT = 'http://schemas.openxmlformats.org/package/2006/content-types'
            override = etree.SubElement(content_types_root, f'{{{CT}}}Override')
            override.set('PartName', f'/word/{footer_name}')
            override.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml')

            fr = etree.Element(f'{{{W}}}footerReference')
            fr.set(f'{{{W}}}type', 'default')
            fr.set(f'{{{R}}}id', rid)
            sectPr.insert(0, fr)
            return footer_name

        footer_targets = {}
        toc_footer_target = None
        if toc_section and toc_section <= len(section_props):
            toc_sect = section_props[toc_section - 1]
            ensure_pgnum(toc_sect, 'upperRoman', 1)
            _, target = default_footer_target(toc_sect)
            if not target:
                target = add_footer_part(toc_sect, ' PAGE \\* ROMAN ')
            toc_footer_target = target
            footer_targets[target] = ' PAGE \\* ROMAN '

        if body_section and body_section <= len(section_props):
            body_sect = section_props[body_section - 1]
            ensure_pgnum(body_sect, 'decimal', 1)
            fr, target = default_footer_target(body_sect)
            if not target or target == toc_footer_target:
                target = add_footer_part(body_sect, ' PAGE \\# "000" ')
            footer_targets[target] = ' PAGE \\# "000" '

        def patch_footer(data, instr):
            footer = etree.fromstring(data)
            patched = False
            for node in footer.findall('.//' + f'{{{W}}}instrText'):
                if node.text and 'PAGE' in node.text:
                    node.text = instr
                    patched = True
            if patched:
                for node in footer.findall('.//' + f'{{{W}}}t'):
                    if node.text and node.text.strip().isdigit():
                        node.text = 'I' if 'ROMAN' in instr else '001'
                        break
            return etree.tostring(footer, xml_declaration=True, encoding='UTF-8', standalone=True)

        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
                elif item.filename == 'word/_rels/document.xml.rels':
                    data = etree.tostring(rels_root, xml_declaration=True, encoding='UTF-8', standalone=True)
                elif item.filename == '[Content_Types].xml':
                    data = etree.tostring(content_types_root, xml_declaration=True, encoding='UTF-8', standalone=True)
                elif item.filename.startswith('word/footer'):
                    rel_target = item.filename.replace('word/', '')
                    if rel_target in footer_targets:
                        data = patch_footer(data, footer_targets[rel_target])
                zout.writestr(item, data)
            for name, data in new_parts.items():
                zout.writestr(name, data)

    os.replace(tmp_path, docx_path)


def _load_project_config(input_path):
    """Load book.toml from the project root or near the manuscript."""
    input_dir = Path(input_path).resolve().parent
    candidates = [
        Path.cwd() / 'book.toml',
        input_dir / 'book.toml',
        input_dir.parent / 'book.toml',
    ]
    config_path = next((p for p in candidates if p.is_file()), None)
    if config_path is None:
        return {}, None
    with config_path.open('rb') as f:
        return tomllib.load(f), config_path


def _safe_filename_part(value, fallback='待定书名'):
    """Remove characters Windows does not allow in file names."""
    value = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '', str(value)).strip().rstrip('.')
    return value or fallback


def _format_filename(pattern, title, author, prefix):
    now = datetime.now()
    replacements = {
        'prefix': _safe_filename_part(prefix, ''),
        'title': _safe_filename_part(title),
        'author': _safe_filename_part(author, '未署名作者'),
        'yyyyMMdd': now.strftime('%Y%m%d'),
        'MMddHHmm': now.strftime('%m%d%H%M'),
        'HHmm': now.strftime('%H%M'),
        'yyyy': now.strftime('%Y'),
        'MM': now.strftime('%m'),
        'dd': now.strftime('%d'),
    }
    filename = pattern
    for key, value in replacements.items():
        filename = filename.replace(f'{{{key}}}', value)
    filename = _safe_filename_part(filename)
    if not filename.lower().endswith('.docx'):
        filename += '.docx'
    return filename


def _resolve_output_path(output_arg, cfg, title, author):
    export_cfg = cfg.get('bookcraft', {}).get('export', {})
    default_dir = export_cfg.get('output_directory', 'exports')
    prefix = export_cfg.get('filename_prefix', '')
    pattern = export_cfg.get('filename_pattern', '{title}-{author}-{yyyyMMdd}-{HHmm}.docx')

    if output_arg:
        supplied = Path(output_arg)
        output_dir = supplied.parent if supplied.suffix.lower() == '.docx' else supplied
    else:
        output_dir = Path(default_dir)

    filename = _format_filename(pattern, title, author, prefix)
    output_dir.mkdir(parents=True, exist_ok=True)
    return str(output_dir / filename)


def _validate_subtitle(subtitle, cfg):
    rules = cfg.get('bookcraft', {}).get('content_rules', {}).get('book_title', {})
    max_chars = int(rules.get('subtitle_max_chars', 15))
    separators = rules.get('subtitle_forbidden_separators', ['：', ':', '——', '--', '｜', '|'])
    if len(subtitle) > max_chars:
        raise ValueError(f'副书名不能超过 {max_chars} 个字符，当前为 {len(subtitle)} 个字符。')
    found = next((sep for sep in separators if sep and sep in subtitle), None)
    if found:
        raise ValueError(f'副书名必须是一句完整短句，不能使用“{found}”拆成两部分。')


def _validate_manuscript_structure(input_path, cfg):
    """Block export when required references or optional back matter are empty."""
    text = Path(input_path).read_text(encoding='utf-8-sig')
    sections = []
    current_title = None
    current_lines = []

    for line in text.splitlines():
        heading = re.match(r'^#\s+(.+?)\s*$', line)
        if heading:
            if current_title is not None:
                sections.append((current_title, current_lines))
            current_title = re.sub(r'^\d+\s+', '', heading.group(1).strip())
            current_lines = []
        elif current_title is not None:
            current_lines.append(line)
    if current_title is not None:
        sections.append((current_title, current_lines))

    placeholders = {
        '待补充', '暂无', '无', '待定', '待完善', '待核查',
        '待补充参考文献', '暂无参考文献'
    }

    def meaningful_lines(lines):
        result = []
        for raw in lines:
            value = raw.strip()
            if not value or value.startswith('<!--'):
                continue
            normalized = re.sub(r'^[\-*+\d.、)）\s]+', '', value).strip('。；; ')
            if normalized in placeholders:
                continue
            result.append(value)
        return result

    references = [lines for title, lines in sections if title == '参考文献']
    rules = cfg.get('bookcraft', {}).get('content_rules', {}).get('references', {})
    min_entries = int(rules.get('min_entries', 1))
    if not references:
        raise ValueError('书稿缺少必选的“参考文献”一级标题。')
    reference_entries = meaningful_lines(references[-1])
    if len(reference_entries) < min_entries:
        raise ValueError(f'参考文献必须至少包含 {min_entries} 条真实有效文献，不能使用占位文字。')

    optional_back_matter = {'结语', '后记', '致谢', '附录', '术语表', '作者简介'}
    for section_title, lines in sections:
        if section_title in optional_back_matter and not meaningful_lines(lines):
            raise ValueError(
                f'可选后置页“{section_title}”没有真实内容，应从 manuscript.md 中删除后再导出。'
            )
    return len(reference_entries)


def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown manuscript to Word (.docx) with classic Chinese book layout"
    )
    parser.add_argument("--input", required=True, help="Input Markdown file (manuscript.md)")
    parser.add_argument("--output", default=None,
                        help="Output directory or legacy .docx path; filename is generated automatically")
    parser.add_argument("--title", default="", help="Book title (for reference template)")
    parser.add_argument("--subtitle", default="", help="Book subtitle")
    parser.add_argument("--author", default="", help="Author name")
    parser.add_argument("--reference-docx", default=None,
                        help="Path to existing reference.docx (generated if not provided)")
    args = parser.parse_args()

    # Validate input
    if not os.path.isfile(args.input):
        print(f"ERROR: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    cfg, config_path = _load_project_config(args.input)
    std_book = cfg.get('book', {})
    bookcraft = cfg.get('bookcraft', {})
    bookcraft_book = bookcraft.get('book', {})

    title = args.title.strip() or str(std_book.get('title', '')).strip() or '待定书名'
    subtitle = args.subtitle.strip() or str(bookcraft_book.get('subtitle', '')).strip()
    authors = std_book.get('authors') or []
    author = args.author.strip() or ('、'.join(str(item) for item in authors if item) if authors else '') or '未署名作者'

    try:
        _validate_subtitle(subtitle, cfg)
        _validate_manuscript_structure(args.input, cfg)
    except ValueError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        sys.exit(2)

    output_path = _resolve_output_path(args.output, cfg, title, author)
    if config_path:
        print(f"Using project config: {config_path}", file=sys.stderr)

    # Ensure dependencies
    ensure_pandoc()
    ensure_python_docx()

    # Generate or locate reference template
    if args.reference_docx and os.path.isfile(args.reference_docx):
        ref_docx = args.reference_docx
        print(f"Using existing reference template: {ref_docx}", file=sys.stderr)
    else:
        ref_docx = os.path.join(os.path.dirname(output_path) or '.', 'reference.docx')
        create_reference_docx(ref_docx, title, subtitle, author)

    # Convert
    pandoc_convert(args.input, output_path, ref_docx)

    # Post-process
    post_process_docx(output_path, title, subtitle, author)
    update_toc_page_numbers_with_word(output_path)

    # Report
    file_size = os.path.getsize(output_path)
    print(f"\n✓ Word document generated: {output_path}", file=sys.stderr)
    print(f"  Size: {file_size / 1024:.1f} KB", file=sys.stderr)


if __name__ == "__main__":
    main()
